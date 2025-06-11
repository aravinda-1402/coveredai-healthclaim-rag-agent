import os
from flask import Flask, request, jsonify, send_from_directory, session, redirect, url_for
from flask_cors import CORS
from flask_dance.contrib.google import make_google_blueprint, google
from flask_session import Session
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.chains.summarize import load_summarize_chain
from langchain.docstore.document import Document
from dotenv import load_dotenv
import re
import uuid
import hashlib
from datetime import datetime, timedelta
import json
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import pandas as pd
import PyPDF2
from functools import wraps
import docx
import time
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1' 

# Load environment variables
load_dotenv()

app = Flask(__name__, static_folder='../frontend/build')
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'your-secret-key-here')

# Configure Flask-Session
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = './flask_session'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)
Session(app)

# Configure CORS with credentials support
CORS(app, supports_credentials=True, resources={
    r"/*": {
        "origins": ["http://localhost:3000"],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"],
        "supports_credentials": True
    }
})

# Configure Google OAuth
blueprint = make_google_blueprint(
    client_id=os.getenv('GOOGLE_CLIENT_ID'),
    client_secret=os.getenv('GOOGLE_CLIENT_SECRET'),
    scope=["https://www.googleapis.com/auth/userinfo.email",
    "https://www.googleapis.com/auth/userinfo.profile",
    "openid"],
    redirect_to="google_authorized"
)
app.register_blueprint(blueprint, url_prefix="/login")

# Configure upload settings
UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'docx'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Initialize OpenAI and embeddings
llm = ChatOpenAI(
    model_name="gpt-3.5-turbo",
    temperature=0,
)
embeddings = OpenAIEmbeddings()

# Store document vectors and conversations per user
user_data = {}

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not google.authorized:
            return jsonify({'error': 'Not authenticated'}), 401
        return f(*args, **kwargs)
    return decorated_function

def get_user_info():
    """Get user info from Google"""
    if not google.authorized:
        return None
    
    resp = google.get("/oauth2/v1/userinfo")
    if resp.ok:
        return resp.json()
    return None

@app.route("/google/authorized")
def google_authorized():
    if not google.authorized:
        return redirect(url_for("google.login"))
    
    user_info = get_user_info()
    if user_info:
        session['user_id'] = user_info['email']
        session['user_name'] = user_info['name']
        session['user_email'] = user_info['email']
        session.permanent = True
        
        # Initialize user data if needed
        if user_info['email'] not in user_data:
            user_data[user_info['email']] = {
                'document_vectors': {},
                'conversations': [],
                'summaries': {}
            }
    
    return redirect("http://localhost:3000")  # Redirect to frontend

@app.route("/auth/user")
def get_current_user():
    """Get current user info"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
    
    return jsonify({
        'name': session.get('user_name'),
        'email': session.get('user_email')
    })

@app.route("/auth/logout")
def logout():
    """Logout user"""
    session.clear()
    return jsonify({'message': 'Logged out successfully'})

def get_user_id():
    """Get or create user ID from session"""
    if 'user_id' not in session:
        session['user_id'] = str(datetime.now().timestamp())  # Create new user ID
    return session['user_id']

def init_user_data(user_id):
    """Initialize user data structure if not exists"""
    if user_id not in user_data:
        user_data[user_id] = {
            'document_vectors': {},
            'conversations': [],
            'summaries': {}
        }

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def sanitize_text(text):
    """Enhanced HIPAA-compliant PII/PHI sanitization"""
    patterns = {
        'ssn': r'\b\d{3}[-.]?\d{2}[-.]?\d{4}\b',
        'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'phone': r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
        'dob': r'\b(0[1-9]|1[0-2])[-/](0[1-9]|[12]\d|3[01])[-/](\d{2}|\d{4})\b',
        'name': r'\b(?:Mr\.|Mrs\.|Ms\.|Dr\.|Prof\.) [A-Z][a-z]+ [A-Z][a-z]+\b',
        'address': r'\b\d{1,5} [A-Za-z\s]{1,30}(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd)\b',
        'mrn': r'\b(?:MRN|Medical Record Number):\s*\d{6,}\b'
    }
    
    for phi_type, pattern in patterns.items():
        text = re.sub(pattern, f'[{phi_type.upper()}]', text)
    
    return text

def extract_benefits(text):
    """Extract key insurance benefits and format as a table"""
    prompt = """Extract the following insurance benefits from the text and format as a JSON object with this exact structure:
    {
        "deductible": {
            "individual": "amount or 'Not specified'",
            "family": "amount or 'Not specified'"
        },
        "out-of-pocket maximum": {
            "individual": "amount or 'Not specified'",
            "family": "amount or 'Not specified'"
        },
        "primary care copay": "amount or 'Not specified'",
        "specialist copay": "amount or 'Not specified'",
        "emergency room copay": "amount or 'Not specified'",
        "urgent care copay": "amount or 'Not specified'",
        "prescription drug coverage": "description or 'Not specified'",
        "mental health copay": "amount or 'Not specified'"
    }
    
    Extract exact amounts when available. Include dollar signs and any relevant notes about limitations."""
    
    response = llm.invoke(f"{prompt}\n\nText: {text}")
    try:
        benefits = json.loads(response.content)
    except:
        # Fallback default structure if parsing fails
        benefits = {
            "deductible": {
                "individual": "Not specified",
                "family": "Not specified"
            },
            "out-of-pocket maximum": {
                "individual": "Not specified",
                "family": "Not specified"
            },
            "primary care copay": "Not specified",
            "specialist copay": "Not specified",
            "emergency room copay": "Not specified",
            "urgent care copay": "Not specified",
            "prescription drug coverage": "Not specified",
            "mental health copay": "Not specified"
        }
    
    return benefits

def generate_suggested_questions(text):
    """Generate relevant questions based on the document content"""
    prompt = """Based on this insurance document, generate 5 common questions that a person might ask. 
    Return them as a JSON array of strings. Make them specific to the actual content."""
    
    response = llm.invoke(f"{prompt}\n\nText: {text}")
    try:
        questions = json.loads(response.content)
    except:
        questions = ["What is my deductible?",
                    "What is my out-of-pocket maximum?",
                    "What is my copay for specialist visits?",
                    "What is covered under preventive care?",
                    "How much do I pay for prescription drugs?"]
    
    return questions

def process_pdf(file_path):
    """Extract text from PDF file"""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n\n"
        return text
    except Exception as e:
        print(f"Error processing PDF: {str(e)}")
        raise

def process_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        print(f"Error processing DOCX: {str(e)}")
        raise

def chunk_text(text, chunk_size=500, chunk_overlap=50):
    """Split text into smaller chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return text_splitter.split_text(text)

def generate_pdf_report(conversation, filename):
    """Generate a PDF report of the conversation"""
    report_path = os.path.join(app.config['UPLOAD_FOLDER'], f'report_{filename}.pdf')
    doc = SimpleDocTemplate(report_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30
    )
    story.append(Paragraph("Insurance Q&A Report", title_style))
    story.append(Spacer(1, 12))
    
    # Add timestamp
    story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Add conversations
    for entry in conversation:
        # Question
        story.append(Paragraph(f"Q: {entry['question']}", styles['Heading3']))
        story.append(Spacer(1, 6))
        
        # Answer
        story.append(Paragraph(f"A: {entry['answer']}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Sources
        if entry.get('sources'):
            story.append(Paragraph("Sources:", styles['Heading4']))
            for source in entry['sources']:
                story.append(Paragraph(f"• {source}", styles['Normal']))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    return report_path

def validate_insurance_document(text):
    """
    Validate if the document is insurance-related using LLM.
    Returns (is_valid, reason)
    """
    try:
        validation_prompt = """
        Analyze this document text and determine if it is specifically a health insurance or medical insurance document.
        
        Required Elements (document must contain at least 3 of these):
        1. Health insurance policy details
        2. Medical coverage information
        3. Healthcare provider networks
        4. Medical procedures or treatments
        5. Insurance premiums or payments
        6. Health plan benefits
        7. Medical claims information
        
        Specific Indicators to Look For:
        - Insurance policy numbers or group numbers
        - Healthcare provider or insurance company names
        - Medical procedure codes (CPT, ICD)
        - Health insurance terminology (deductible, copay, coinsurance)
        - Medical benefit descriptions
        - Network coverage details (in-network, out-of-network)
        - Healthcare service categories
        - Prescription drug coverage
        - Medical claims processing information
        - Member/patient information fields
        
        Return a JSON object with the following fields:
        {
            "is_insurance": boolean,  // true only if it's a health/medical insurance document
            "document_type": string,  // e.g. "Health Insurance Policy", "Medical Claim Form", "Benefits Statement", etc.
            "confidence_score": number,  // 0-100 based on how many required elements and indicators are found
            "found_elements": [string],  // list of required elements found in the document
            "found_indicators": [string],  // list of specific indicators found
            "reason": string  // explanation of the determination
        }
        
        If it's not a health/medical insurance document, explain what type of document it appears to be instead.
        
        Text to analyze:
        """
        
        response = llm.invoke(validation_prompt + "\n\n" + text[:2000])  # Use first 2000 chars for better assessment
        result = json.loads(str(response.content))
        
        is_insurance = result.get("is_insurance", False)
        confidence_score = result.get("confidence_score", 0)
        
        # Only consider it valid if it's a health insurance document with high confidence
        is_valid = is_insurance and confidence_score >= 70
        
        # Construct detailed reason
        if is_valid:
            reason = f"This appears to be a {result.get('document_type', 'health insurance document')}. "
            if result.get("found_elements"):
                reason += f"Found key elements: {', '.join(result['found_elements'][:3])}. "
            if result.get("found_indicators"):
                reason += f"Identified specific insurance indicators: {', '.join(result['found_indicators'][:3])}."
        else:
            if not is_insurance:
                reason = f"This appears to be {result.get('reason', 'not a health insurance document')}."
            else:
                reason = f"This document has low confidence ({confidence_score}%) of being a valid health insurance document. "
                reason += result.get('reason', '')
        
        return is_valid, reason
        
    except Exception as e:
        print(f"Error in document validation: {str(e)}")
        return False, "Error during validation"

def is_insurance_document(text, min_confidence=3):
    """
    Validate if a document is an insurance document by checking for key insurance-related terms and patterns.
    Returns True if the document appears to be an insurance document, False otherwise.
    """
    # Key insurance terms and phrases to look for
    insurance_indicators = [
        r'\b(health|medical)\s+insurance\b',
        r'\b(plan|policy|coverage)\s+(summary|details|information)\b',
        r'\bbenefits?\s+(summary|overview|details)\b',
        r'\b(in|out)\s*-?\s*of\s*-?\s*network\b',
        r'\bdeductible\b',
        r'\bcopay(ment)?\b',
        r'\bcoinsurance\b',
        r'\b(covered|eligible)\s+services\b',
        r'\b(prescription|rx)\s+drugs?\b',
        r'\b(prior\s+)?authorization\b',
        r'\bpreventive\s+care\b',
        r'\bemergency\s+(room|services)\b',
        r'\bout\s*-?\s*of\s*-?\s*pocket\s+(maximum|limit)\b',
        r'\bpremium\b',
        r'\bprovider\s+network\b',
        r'\bclaims?\b',
        r'\benrollment\b',
        r'\beligibility\b'
    ]
    
    # Count how many insurance indicators are found
    confidence_score = 0
    found_terms = set()
    
    # Take first 2000 characters for analysis to avoid processing entire large documents
    sample_text = text[:2000].lower()
    
    for pattern in insurance_indicators:
        matches = re.findall(pattern, sample_text, re.IGNORECASE)
        if matches:
            confidence_score += 1
            found_terms.add(pattern)
    
    print(f"Insurance document validation - Score: {confidence_score}, Found terms: {found_terms}")
    
    # Return True if we found at least min_confidence number of insurance terms
    return confidence_score >= min_confidence

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    file_ext = file.filename.lower()
    if not (file_ext.endswith('.pdf') or file_ext.endswith('.docx')):
        return jsonify({'error': 'Only PDF and DOCX files are supported'}), 400

    try:
        user_id = get_user_id()
        init_user_data(user_id)  # Initialize user data structure

        # Save the file temporarily
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            # Process the file based on its type
            if file_ext.endswith('.pdf'):
                full_text = process_pdf(file_path)
            else:  # .docx
                full_text = process_docx(file_path)
            
            # Validate if it's an insurance document
            is_insurance, reason = validate_insurance_document(full_text)
            
            if not is_insurance:
                # Delete the temporary file if it's not insurance-related
                os.remove(file_path)
                return jsonify({
                    'error': 'Insurance Document Required',
                    'message': f'This tool is specifically designed for insurance documents only. The uploaded file appears to be about something else. {reason} Please upload a health insurance policy, claims document, or benefits statement.'
                }), 400

            # Split text into smaller chunks
            text_chunks = chunk_text(full_text)
            
            # Store document data
            user_data[user_id]['document_vectors'][filename] = {
                'full_text': full_text,
                'chunks': text_chunks
            }

            # Generate suggested questions only
            try:
                questions_prompt = "Generate 3-5 relevant questions someone might ask about this insurance document. Return as a JSON array of strings."
                questions_response = llm.invoke(questions_prompt + "\n\n" + "\n".join(text_chunks[:2]))
                questions_text = str(questions_response.content) if hasattr(questions_response, 'content') else str(questions_response)
                suggested_questions = json.loads(questions_text) if questions_text else []
            except:
                suggested_questions = []

            return jsonify({
                'message': 'File uploaded successfully',
                'filename': filename,
                'suggested_questions': suggested_questions
            })

        except Exception as e:
            # Clean up the file in case of processing error
            if os.path.exists(file_path):
                os.remove(file_path)
            raise e

    except Exception as e:
        print(f"Error in upload_file: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/ask', methods=['POST'])
@login_required
def ask_question():
    user_id = get_user_id()
    init_user_data(user_id)  # Initialize user data structure
    
    data = request.json
    if not data or 'question' not in data or 'filename' not in data:
        return jsonify({'error': 'Missing question or filename'}), 400
    
    filename = data['filename']
    if filename not in user_data[user_id]['document_vectors']:
        return jsonify({'error': 'Selected document not found'}), 404
    
    try:
        # Get relevant chunks for the question from the selected document only
        question = data['question']
        doc_data = user_data[user_id]['document_vectors'][filename]
        chunks = doc_data['chunks']
        
        if not chunks:
            return jsonify({'error': 'No content found in the selected document'}), 404
        
        # Create documents for vectorstore
        documents = [
            Document(
                page_content=chunk,
                metadata={'source': filename, 'chunk_id': i}
            ) for i, chunk in enumerate(chunks)
        ]
        
        # Create vectorstore with chunks from selected document only
        vectorstore = FAISS.from_documents(documents, embeddings)
        
        # Get most relevant chunks
        relevant_chunks = vectorstore.similarity_search(
            question,
            k=3  # Limit to top 3 most relevant chunks
        )
        
        # Create context from relevant chunks
        context = "\n\n".join([doc.page_content for doc in relevant_chunks])
        
        # Get answer using relevant context only
        prompt = f"""Based on the following insurance document excerpts, answer this question: {question}

Context:
{context}

Answer:"""
        
        answer = llm.invoke(prompt)
        answer_text = str(answer.content) if hasattr(answer, 'content') else str(answer)
        
        # Extract sources
        sources = [{
            'text': doc.page_content,
            'document': doc.metadata['source']
        } for doc in relevant_chunks]
        
        # Store conversation
        user_data[user_id]['conversations'].append({
            'question': question,
            'answer': answer_text,
            'sources': sources,
            'timestamp': datetime.now().isoformat()
        })
        
        return jsonify({
            'answer': answer_text,
            'sources': sources
        })
    
    except Exception as e:
        print(f"Error in ask_question: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/delete-file', methods=['POST'])
@login_required
def delete_file():
    """Delete a specific file for the user"""
    user_id = get_user_id()
    data = request.json
    
    if not data or 'filename' not in data:
        return jsonify({'error': 'No filename provided'}), 400
    
    filename = data['filename']
    
    try:
        # Remove from user data first to prevent new operations
        if filename in user_data[user_id]['document_vectors']:
            # Store file info before deleting from user_data
            file_info = user_data[user_id]['document_vectors'][filename]
            del user_data[user_id]['document_vectors'][filename]
            
            # Delete physical file with retries
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            max_retries = 3
            retry_delay = 1  # seconds
            
            for attempt in range(max_retries):
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    
                    # Clean up any potential temporary files
                    temp_file = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{filename}")
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        
                    # Clean up any potential temporary PDF files for DOCX conversions
                    if filename.lower().endswith('.docx'):
                        pdf_temp = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.splitext(filename)[0]}.pdf")
                        if os.path.exists(pdf_temp):
                            os.remove(pdf_temp)
                    
                    return jsonify({'message': 'File deleted successfully'})
                    
                except PermissionError:
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        continue
                    raise
                except FileNotFoundError:
                    # If file is already deleted, consider it a success
                    return jsonify({'message': 'File deleted successfully'})
                except Exception as e:
                    # Restore user_data if file deletion fails
                    user_data[user_id]['document_vectors'][filename] = file_info
                    raise
        else:
            return jsonify({'error': 'File not found'}), 404
            
    except Exception as e:
        print(f"Error deleting file: {str(e)}")
        return jsonify({'error': f'Failed to delete file: {str(e)}'}), 500

@app.route('/summarize', methods=['POST'])
@login_required
def summarize_plan():
    user_id = get_user_id()
    init_user_data(user_id)  # Initialize user data structure
    
    data = request.json
    if not data or 'filename' not in data:
        return jsonify({'error': 'No filename provided'}), 400
        
    filename = data['filename']
    if filename not in user_data[user_id]['document_vectors']:
        return jsonify({'error': 'File not found'}), 404
        
    try:
        full_text = user_data[user_id]['document_vectors'][filename]['full_text']
        
        # Generate benefits using LLM
        benefits_prompt = """
        Analyze this insurance document and extract key benefits information.
        Return a JSON object with the following structure:
        {
            "deductible": "Amount or 'Not specified'",
            "outOfPocketMax": "Amount or 'Not specified'",
            "coverageDetails": ["List of key coverage points"],
            "copaysAndCoinsurance": {
                "Emergency Room": "Cost or percentage",
                "Inpatient Hospitalization": "Cost or percentage",
                "Mental Health Counseling": "Cost or percentage",
                "Outpatient Surgery": "Cost or percentage",
                "Primary Care Visits": "Cost or percentage",
                "Specialist Visits": "Cost or percentage"
            }
        }
        
        Only include fields where information is explicitly found in the document.
        For any cost/percentage, format it clearly (e.g., "$30 copay" or "20% coinsurance").
        If a value is not found in the document, omit that field entirely.
        """
        
        benefits_response = llm.invoke(benefits_prompt + "\n\n" + full_text)
        benefits = json.loads(str(benefits_response.content))
        
        # Clean up the benefits object to remove any null or undefined values
        if 'copaysAndCoinsurance' in benefits:
            benefits['copaysAndCoinsurance'] = {
                k: v for k, v in benefits['copaysAndCoinsurance'].items()
                if v and v.lower() not in ['not specified', 'not found', 'null', 'undefined', 'none']
            }
            if not benefits['copaysAndCoinsurance']:
                del benefits['copaysAndCoinsurance']
                
        if 'coverageDetails' in benefits:
            benefits['coverageDetails'] = [
                detail for detail in benefits['coverageDetails']
                if detail and detail.lower() not in ['not specified', 'not found', 'null', 'undefined', 'none']
            ]
            if not benefits['coverageDetails']:
                del benefits['coverageDetails']
                
        for key in ['deductible', 'outOfPocketMax']:
            if key in benefits and (not benefits[key] or 
                                  benefits[key].lower() in ['not specified', 'not found', 'null', 'undefined', 'none']):
                del benefits[key]

        # Generate a human-readable summary
        summary_prompt = """
        Create a clear, concise summary of this insurance plan. Focus on:
        1. The most important coverage details
        2. Key benefits and limitations
        3. Notable features or provisions
        
        Make it easy to understand for someone not familiar with insurance terminology.
        Keep it to 2-3 paragraphs maximum.
        """
        
        summary_response = llm.invoke(summary_prompt + "\n\n" + full_text)
        summary = str(summary_response.content)
        
        return jsonify({
            'benefits': benefits,
            'summary': summary
        })
        
    except Exception as e:
        print(f"Error in summarize_plan: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/uploads/<path:filename>')
def download_file(filename):
    """Serve uploaded files"""
    try:
        return send_from_directory(
            app.config['UPLOAD_FOLDER'],
            filename,
            as_attachment=True
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@app.route('/export-report', methods=['POST'])
@login_required
def export_report():
    """Generate and export a PDF report of the Q&A session"""
    user_id = get_user_id()
    
    if not user_data[user_id]['conversations']:
        return jsonify({'error': 'No Q&A history to export'}), 404

    try:
        # Generate unique filename
        report_filename = f"qa_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        report_path = os.path.join(app.config['UPLOAD_FOLDER'], report_filename)
        
        # Create PDF
        doc = SimpleDocTemplate(
            report_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor='#2563eb'  # Blue color
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=10,
            textColor='#1f2937'  # Dark gray
        )
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            spaceBefore=6,
            spaceAfter=12,
            textColor='#374151'  # Medium gray
        )
        source_style = ParagraphStyle(
            'CustomSource',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20,
            textColor='#6b7280',  # Light gray
            spaceBefore=6,
            spaceAfter=12
        )
        
        # Build document content
        story = []
        
        # Title
        story.append(Paragraph("Insurance Q&A Report", title_style))
        story.append(Spacer(1, 12))
        
        # Add timestamp
        story.append(Paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            normal_style
        ))
        story.append(Spacer(1, 20))
        
        # Add user info
        if 'user_name' in session and 'user_email' in session:
            story.append(Paragraph(
                f"User: {session['user_name']} ({session['user_email']})",
                normal_style
            ))
            story.append(Spacer(1, 20))
        
        # Add conversations
        for i, entry in enumerate(user_data[user_id]['conversations'], 1):
            # Question section
            story.append(Paragraph(
                f"Question {i}:",
                heading_style
            ))
            story.append(Paragraph(
                entry['question'],
                normal_style
            ))
            
            # Answer section
            story.append(Paragraph(
                "Answer:",
                heading_style
            ))
            story.append(Paragraph(
                entry['answer'],
                normal_style
            ))
            
            # Sources section
            if entry.get('sources'):
                story.append(Paragraph(
                    "Sources:",
                    heading_style
                ))
                for source in entry['sources']:
                    story.append(Paragraph(
                        f"From {source['document']}:",
                        source_style
                    ))
                    story.append(Paragraph(
                        source['text'],
                        source_style
                    ))
            
            story.append(Spacer(1, 20))
            
            # Add a divider between QA pairs (except for the last one)
            if i < len(user_data[user_id]['conversations']):
                story.append(Paragraph(
                    "_" * 50,  # Simple line divider
                    normal_style
                ))
                story.append(Spacer(1, 20))
        
        # Build the PDF
        doc.build(story)
        
        return jsonify({
            'message': 'Report generated successfully',
            'filename': report_filename
        })
    
    except Exception as e:
        print(f"Error in export_report: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/get-user-files', methods=['GET'])
@login_required
def get_user_files():
    """Get list of files uploaded by the user"""
    user_id = get_user_id()
    init_user_data(user_id)  # Initialize user data structure
    
    files = list(user_data[user_id]['document_vectors'].keys())
    return jsonify({'files': files})

@app.route('/api/compare-plans', methods=['POST'])
@login_required
def compare_plans():
    """Compare multiple insurance plans"""
    user_id = get_user_id()
    init_user_data(user_id)

    if 'file0' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    results = []
    temp_files = []

    try:
        # Process each file
        file_index = 0
        while f'file{file_index}' in request.files and f'label{file_index}' in request.form:
            file = request.files[f'file{file_index}']
            label = request.form[f'label{file_index}']
            
            if file and allowed_file(file.filename):
                # Save and process the file
                filename = secure_filename(file.filename)
                temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(temp_path)
                temp_files.append(temp_path)

                # Convert DOCX to PDF if needed
                if filename.lower().endswith('.docx'):
                    pdf_path = convert_docx_to_pdf(temp_path)
                    temp_files.append(pdf_path)
                    temp_path = pdf_path

                # Extract text and validate
                text = extract_text_from_pdf(temp_path)
                if not is_insurance_document(text):
                    return jsonify({'error': f'File {label} is not a valid insurance document'}), 400

                # Process the document - but don't store in user_data
                doc_data = process_document(text)

                # Extract benefits summary
                summary = extract_benefits_summary(text)
                results.append({
                    'label': label,
                    'filename': filename,
                    'summary': {
                        'deductibles': {
                            'individual': summary.get('individual_deductible', ''),
                            'family': summary.get('family_deductible', '')
                        },
                        'outOfPocketMax': {
                            'individual': summary.get('individual_out_of_pocket_max', ''),
                            'family': summary.get('family_out_of_pocket_max', '')
                        },
                        'copays': {
                            'primaryCare': summary.get('primary_care_copay', ''),
                            'specialist': summary.get('specialist_copay', ''),
                            'emergencyRoom': summary.get('emergency_room_copay', ''),
                            'urgentCare': summary.get('urgent_care_copay', '')
                        },
                        'prescriptionCoverage': summary.get('prescription_coverage', ''),
                        'mentalHealthCoverage': summary.get('mental_health_coverage', '')
                    }
                })

            file_index += 1

        return jsonify(results)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

    finally:
        # Clean up temporary files
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception as e:
                print(f"Error removing temporary file {temp_file}: {str(e)}")

# Serve React frontend
@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def serve(path):
    if path != "" and os.path.exists(app.static_folder + '/' + path):
        return send_from_directory(app.static_folder, path)
    else:
        return send_from_directory(app.static_folder, 'index.html')

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file."""
    try:
        text = ""
        with fitz.open(pdf_path) as doc:
            for page in doc:
                text += page.get_text()
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return ""

def extract_benefits_summary(text):
    """Extract benefits summary from the document text."""
    try:
        # Use the existing extract_benefits function and map its output to our desired format
        benefits = extract_benefits(text)
        
        if isinstance(benefits, str):
            # If extract_benefits returned a string (error), create empty structure
            return {
                'individual_deductible': 'Not found',
                'family_deductible': 'Not found',
                'individual_out_of_pocket_max': 'Not found',
                'family_out_of_pocket_max': 'Not found',
                'primary_care_copay': 'Not found',
                'specialist_copay': 'Not found',
                'emergency_room_copay': 'Not found',
                'urgent_care_copay': 'Not found',
                'prescription_coverage': 'Not found',
                'mental_health_coverage': 'Not found'
            }
        
        # Map the benefits to our desired format
        return {
            'individual_deductible': benefits['deductible']['individual'],
            'family_deductible': benefits['deductible']['family'],
            'individual_out_of_pocket_max': benefits['out-of-pocket maximum']['individual'],
            'family_out_of_pocket_max': benefits['out-of-pocket maximum']['family'],
            'primary_care_copay': benefits['primary care copay'],
            'specialist_copay': benefits['specialist copay'],
            'emergency_room_copay': benefits['emergency room copay'],
            'urgent_care_copay': benefits['urgent care copay'],
            'prescription_coverage': benefits['prescription drug coverage'],
            'mental_health_coverage': benefits['mental health copay']
        }
    except Exception as e:
        print(f"Error extracting benefits summary: {str(e)}")
        return {
            'individual_deductible': 'Error',
            'family_deductible': 'Error',
            'individual_out_of_pocket_max': 'Error',
            'family_out_of_pocket_max': 'Error',
            'primary_care_copay': 'Error',
            'specialist_copay': 'Error',
            'emergency_room_copay': 'Error',
            'urgent_care_copay': 'Error',
            'prescription_coverage': 'Error',
            'mental_health_coverage': 'Error'
        }

def process_document(text, chunk_size=500, chunk_overlap=50):
    """Process a document by chunking text and creating embeddings."""
    try:
        # Sanitize the text to remove any PII/PHI
        sanitized_text = sanitize_text(text)
        
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_text(sanitized_text)
        
        # Create document objects for each chunk
        docs = [Document(page_content=chunk) for chunk in chunks]
        
        # Create embeddings for the chunks
        vectorstore = FAISS.from_documents(docs, embeddings)
        
        return {
            'chunks': chunks,
            'embeddings': vectorstore
        }
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return {
            'chunks': [],
            'embeddings': None
        }

if __name__ == '__main__':
    os.makedirs('./flask_session', exist_ok=True)
    app.run(debug=True)
